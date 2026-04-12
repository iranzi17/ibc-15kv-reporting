from __future__ import annotations

import base64
import html
import os
import re
import shutil
import subprocess
import tempfile
import textwrap
import zipfile
from contextlib import nullcontext
from io import BytesIO
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
from sheets import (
    CACHE_FILE,
    append_rows_to_sheet,
    get_sheet_data,
    get_unique_sites_and_dates,
    load_offline_cache,
)
from report import generate_reports, signatories_for_row
from report_structuring import REPORT_HEADERS, clean_and_structure_report
from ui import render_workwatch_header, set_background
from ui_hero import render_hero

st.set_page_config(page_title="WorkWatch - Site Intelligence", layout="wide")

DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
OPENAI_API_KEY_SESSION_KEY = "openai_api_key"
OPENAI_CHAT_MESSAGES_KEY = "openai_chat_messages"
OPENAI_PREVIOUS_RESPONSE_ID_KEY = "openai_previous_response_id"
OPENAI_MODEL_SESSION_KEY = "openai_model"


def _safe_columns(*args, **kwargs):
    """Call st.columns falling back to positional-only call for stubs."""
    columns_fn = getattr(st, "columns", None)
    if not callable(columns_fn):
        return (nullcontext(), nullcontext())

    try:
        return columns_fn(*args, **kwargs)
    except TypeError:
        return columns_fn(*args)


def _safe_markdown(markdown: str, **kwargs) -> None:
    """Call st.markdown when available (tests provide a stub without it)."""
    markdown_fn = getattr(st, "markdown", None)
    if callable(markdown_fn):
        markdown_fn(markdown, **kwargs)


def _safe_checkbox(label: str, *, value=False, key=None):
    """Return st.checkbox value or default when unavailable."""
    checkbox_fn = getattr(st, "checkbox", None)
    if callable(checkbox_fn):
        return checkbox_fn(label, value=value, key=key)
    return value


def _safe_caption(text: str) -> None:
    """Call st.caption when available."""
    caption_fn = getattr(st, "caption", None)
    if callable(caption_fn):
        caption_fn(text)


def _safe_data_editor(df: pd.DataFrame, **kwargs) -> pd.DataFrame:
    """Return data editor result, or input frame when editor is unavailable."""
    editor_fn = getattr(st, "data_editor", None)
    if not callable(editor_fn):
        return df

    try:
        edited = editor_fn(df, **kwargs)
    except TypeError:
        edited = editor_fn(df)

    if isinstance(edited, pd.DataFrame):
        return edited
    return df


def _safe_selectbox(label: str, options: list[Any], index: int = 0, **kwargs):
    """Return st.selectbox value, with a deterministic fallback for test stubs."""
    selectbox_fn = getattr(st, "selectbox", None)
    if callable(selectbox_fn):
        try:
            return selectbox_fn(label, options=options, index=index, **kwargs)
        except TypeError:
            return selectbox_fn(label, options, index)

    if not options:
        return None
    safe_index = max(0, min(index, len(options) - 1))
    return options[safe_index]


def _safe_image(images, **kwargs) -> None:
    """Call st.image when available."""
    image_fn = getattr(st, "image", None)
    if callable(image_fn):
        image_fn(images, **kwargs)


def _safe_text_input(label: str, value: str = "", **kwargs) -> str:
    """Call st.text_input when available, falling back to the default value."""
    text_input_fn = getattr(st, "text_input", None)
    if callable(text_input_fn):
        try:
            return text_input_fn(label, value=value, **kwargs)
        except TypeError:
            return text_input_fn(label, value)
    return value


def _safe_write(value: object) -> None:
    """Call st.write when available."""
    write_fn = getattr(st, "write", None)
    if callable(write_fn):
        write_fn(value)


def _safe_expander(label: str, *, expanded: bool = False):
    """Return st.expander context or a nullcontext fallback."""
    expander_fn = getattr(st, "expander", None)
    if callable(expander_fn):
        try:
            return expander_fn(label, expanded=expanded)
        except TypeError:
            return expander_fn(label)
    return nullcontext()


def _safe_chat_message(role: str):
    """Return st.chat_message context or a nullcontext fallback."""
    chat_message_fn = getattr(st, "chat_message", None)
    if callable(chat_message_fn):
        return chat_message_fn(role)
    return nullcontext()


def _safe_chat_input(prompt: str, **kwargs) -> str | None:
    """Return st.chat_input value or None when unavailable."""
    chat_input_fn = getattr(st, "chat_input", None)
    if callable(chat_input_fn):
        try:
            return chat_input_fn(prompt, **kwargs)
        except TypeError:
            return chat_input_fn(prompt)
    return None


def _safe_spinner(text: str):
    """Return st.spinner context or a nullcontext fallback."""
    spinner_fn = getattr(st, "spinner", None)
    if callable(spinner_fn):
        return spinner_fn(text)
    return nullcontext()


def _streamlit_secret(name: str, default: str = "") -> str:
    """Read one Streamlit secret value without failing when secrets are unavailable."""
    try:
        secrets = getattr(st, "secrets", None)
        if secrets is None:
            return default
        if hasattr(secrets, "get"):
            value = secrets.get(name, default)
        elif name in secrets:
            value = secrets[name]
        else:
            value = default
    except Exception:
        return default

    if value is None:
        return default
    return str(value).strip()


def _load_openai_api_key() -> str:
    """Return the active OpenAI API key from session, env, or Streamlit secrets."""
    session_key = str(st.session_state.get(OPENAI_API_KEY_SESSION_KEY, "") or "").strip()
    if session_key:
        return session_key

    env_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if env_key:
        return env_key

    return _streamlit_secret("OPENAI_API_KEY")


def _default_openai_model() -> str:
    """Return the preferred chat model, allowing env or secrets overrides."""
    session_model = str(st.session_state.get(OPENAI_MODEL_SESSION_KEY, "") or "").strip()
    if session_model:
        return session_model

    env_model = os.environ.get("OPENAI_MODEL", "").strip()
    if env_model:
        return env_model

    secret_model = _streamlit_secret("OPENAI_MODEL")
    if secret_model:
        return secret_model

    return DEFAULT_OPENAI_MODEL


def _openai_sdk_ready() -> tuple[bool, str]:
    """Check whether the OpenAI SDK can be imported."""
    try:
        from openai import OpenAI  # noqa: F401
    except ImportError as exc:
        return False, str(exc)
    return True, ""


def _extract_openai_output_text(response) -> str:
    """Read the text output from an OpenAI Responses API object."""
    text = str(getattr(response, "output_text", "") or "").strip()
    if text:
        return text

    fragments: list[str] = []
    for item in getattr(response, "output", []) or []:
        for content in getattr(item, "content", []) or []:
            content_text = str(getattr(content, "text", "") or "").strip()
            if content_text:
                fragments.append(content_text)

    return "\n".join(fragments).strip()


def _request_openai_reply(prompt: str, *, api_key: str, model: str) -> tuple[str, str]:
    """Send one chat turn to OpenAI and return (reply_text, response_id)."""
    from openai import OpenAI

    client = OpenAI(api_key=api_key)
    request_kwargs = {
        "model": model,
        "input": [{"role": "user", "content": prompt}],
    }

    previous_response_id = st.session_state.get(OPENAI_PREVIOUS_RESPONSE_ID_KEY)
    if previous_response_id:
        request_kwargs["previous_response_id"] = previous_response_id

    response = client.responses.create(**request_kwargs)
    reply_text = _extract_openai_output_text(response)
    if not reply_text:
        reply_text = "No text response was returned."

    return reply_text, str(getattr(response, "id", "") or "")


def _clear_openai_chat() -> None:
    """Reset the chat transcript and response threading state."""
    st.session_state[OPENAI_CHAT_MESSAGES_KEY] = []
    st.session_state.pop(OPENAI_PREVIOUS_RESPONSE_ID_KEY, None)


def _render_openai_chat_panel() -> None:
    """Render the ChatGPT-style assistant UI for the Streamlit app."""
    st.subheader("ChatGPT Assistant")
    _safe_caption(
        "This uses the OpenAI API from your app. It does not connect to your personal ChatGPT chat history."
    )

    with _safe_expander("Chat settings", expanded=False):
        entered_key = _safe_text_input(
            "OpenAI API key (session only)",
            value="",
            type="password",
            key="openai_api_key_input",
            placeholder="sk-...",
        ).strip()
        if entered_key:
            st.session_state[OPENAI_API_KEY_SESSION_KEY] = entered_key

        active_key = _load_openai_api_key()
        key_source = "session input"
        if not st.session_state.get(OPENAI_API_KEY_SESSION_KEY):
            if os.environ.get("OPENAI_API_KEY", "").strip():
                key_source = "environment variable"
            elif _streamlit_secret("OPENAI_API_KEY"):
                key_source = "Streamlit secrets"
            else:
                key_source = "not configured"

        model_value = _safe_text_input(
            "OpenAI model",
            value=_default_openai_model(),
            key="openai_model_input",
            placeholder=DEFAULT_OPENAI_MODEL,
        ).strip()
        st.session_state[OPENAI_MODEL_SESSION_KEY] = model_value or DEFAULT_OPENAI_MODEL

        clear_key_col, clear_chat_col = _safe_columns(2, gap="large")
        with clear_key_col:
            if st.button("Forget session API key"):
                st.session_state.pop(OPENAI_API_KEY_SESSION_KEY, None)
                st.session_state["openai_api_key_input"] = ""
        with clear_chat_col:
            if st.button("Clear ChatGPT history"):
                _clear_openai_chat()

        if active_key:
            st.success(f"OpenAI key loaded from {key_source}.")
        else:
            st.info(
                "Add OPENAI_API_KEY to .streamlit/secrets.toml or your environment, or paste it above for this browser session."
            )

    sdk_ready, sdk_error = _openai_sdk_ready()
    if not sdk_ready:
        st.warning(
            "OpenAI SDK is not installed yet. Run `pip install -r requirements.txt` and reload the app. "
            f"Detail: {sdk_error}"
        )
        return

    messages = st.session_state.setdefault(OPENAI_CHAT_MESSAGES_KEY, [])
    prompt = _safe_chat_input("Ask ChatGPT anything about your reports or field work.")

    if prompt:
        api_key = _load_openai_api_key()
        if not api_key:
            st.warning("OpenAI API key is required before you can start chatting.")
        else:
            model = _default_openai_model()
            messages.append({"role": "user", "content": prompt})
            try:
                with _safe_spinner("Waiting for OpenAI..."):
                    reply_text, response_id = _request_openai_reply(
                        prompt,
                        api_key=api_key,
                        model=model,
                    )
            except Exception as exc:
                messages.pop()
                st.error(f"OpenAI request failed: {exc}")
            else:
                messages.append({"role": "assistant", "content": reply_text})
                if response_id:
                    st.session_state[OPENAI_PREVIOUS_RESPONSE_ID_KEY] = response_id

    for message in messages:
        with _safe_chat_message(str(message.get("role", "assistant"))):
            _safe_write(message.get("content", ""))


def _load_sheet_context():
    """Return (data_rows, sites, error) while isolating failures."""
    try:
        rows = get_sheet_data()
        data_rows = rows[1:] if rows else []
        sites, _ = get_unique_sites_and_dates(data_rows)
        return data_rows, list(sites), None
    except Exception as exc:  # pragma: no cover - user notification
        return [], [], exc


def _rows_to_structured_data(rows: list[list[str]]) -> list[dict[str, str]]:
    """Convert row lists into dicts keyed by REPORT_HEADERS names."""
    structured = []
    header_count = len(REPORT_HEADERS)
    for row in rows:
        padded = (row + [""] * header_count)[:header_count]
        entry = {header: value for header, value in zip(REPORT_HEADERS, padded)}
        structured.append(entry)
    return structured


def _normalized_review_rows(df: pd.DataFrame) -> list[list[str]]:
    """Normalize edited table values back into report row lists."""
    if df is None or df.empty:
        return []

    normalized = df.reindex(columns=REPORT_HEADERS).fillna("")
    rows: list[list[str]] = []
    for row in normalized.values.tolist():
        rows.append([str(cell).strip() for cell in row])
    return rows


def _bytes_to_data_uri(image_bytes: bytes, mime: str = "image/jpeg") -> str:
    """Return data URI for raw bytes."""
    return f"data:{mime};base64,{base64.b64encode(image_bytes).decode('utf-8')}"


def _preview_text(value: object) -> str:
    """Escape text for HTML preview blocks."""
    return html.escape(str(value or "")).replace("\n", "<br>")


def _docx_cell_html(cell, related_parts) -> str:
    """Render one DOCX table cell into HTML, including embedded images."""
    parts = []
    for para in cell.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(f"<div>{_preview_text(text)}</div>")

    rel_ids = sorted(set(re.findall(r'r:embed="([^"]+)"', cell._tc.xml)))
    for rel_id in rel_ids:
        image_part = related_parts.get(rel_id)
        if not image_part:
            continue
        mime = getattr(image_part, "content_type", "image/jpeg")
        parts.append(
            f'<img class="tpl-image" src="{_bytes_to_data_uri(image_part.blob, mime)}" alt="Embedded image" />'
        )

    return "".join(parts) if parts else "&nbsp;"


def _render_template_preview_html(docx_bytes: bytes) -> str:
    """Build a template-size preview using generated DOCX structure and section sizes."""
    from docx import Document
    from docx.shared import Mm

    document = Document(BytesIO(docx_bytes))

    emu_per_mm = Mm(1).emu
    section = document.sections[0] if document.sections else None

    def emu_to_mm(value, default: float) -> float:
        try:
            if value is None:
                return default
            return float(value) / float(emu_per_mm)
        except Exception:
            return default

    page_w_mm = emu_to_mm(getattr(section, "page_width", None), 210.0)
    page_h_mm = emu_to_mm(getattr(section, "page_height", None), 297.0)
    margin_left_mm = emu_to_mm(getattr(section, "left_margin", None), 20.0)
    margin_right_mm = emu_to_mm(getattr(section, "right_margin", None), 20.0)
    margin_top_mm = emu_to_mm(getattr(section, "top_margin", None), 20.0)
    margin_bottom_mm = emu_to_mm(getattr(section, "bottom_margin", None), 20.0)

    para_html_parts = []
    for para in document.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            para_html_parts.append(f"<p>{_preview_text(txt)}</p>")
    paragraphs_html = "".join(para_html_parts)

    table_html = '<div class="tpl-muted">No table content found in template preview.</div>'
    if document.tables:
        rows_html = []
        first_table = document.tables[0]
        for row in first_table.rows:
            cells_html = []
            for cell in row.cells:
                cell_html = _docx_cell_html(cell, document.part.related_parts)
                cells_html.append(f"<td>{cell_html}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
        table_html = f"<table class='tpl-table'>{''.join(rows_html)}</table>"

    return textwrap.dedent(
        f"""
        <style>
          .tpl-shell {{
            border: 1px solid #d0d5dd;
            border-radius: 10px;
            background: #f5f7fa;
            padding: 12px;
            overflow-x: auto;
          }}
          .tpl-page {{
            width: min({page_w_mm:.2f}mm, 100%);
            min-height: {page_h_mm:.2f}mm;
            margin: 0 auto;
            box-sizing: border-box;
            padding: {margin_top_mm:.2f}mm {margin_right_mm:.2f}mm {margin_bottom_mm:.2f}mm {margin_left_mm:.2f}mm;
            background: #fff;
            border: 1px solid #c3cad4;
            box-shadow: 0 8px 24px rgba(16, 24, 40, 0.15);
            color: #111827;
            font-family: "Times New Roman", Georgia, serif;
            font-size: 11pt;
            line-height: 1.35;
          }}
          .tpl-page p {{ margin: 0 0 6px 0; }}
          .tpl-table {{
            width: 100%;
            border-collapse: collapse;
            table-layout: fixed;
            margin-top: 8px;
          }}
          .tpl-table td {{
            border: 1px solid #b9c1cc;
            padding: 4px 5px;
            vertical-align: top;
            font-size: 10.5pt;
            word-break: break-word;
          }}
          .tpl-image {{
            display: block;
            max-width: 100%;
            height: auto;
            margin: 3px 0;
          }}
          .tpl-muted {{ color: #667085; font-style: italic; }}
          @media (max-width: 900px) {{
            .tpl-page {{
              width: 100%;
              min-height: auto;
              padding: 14px;
            }}
          }}
        </style>
        <div class="tpl-shell">
          <div class="tpl-page">
            {paragraphs_html}
            {table_html}
          </div>
        </div>
        """
    ).strip()


def _build_single_report_docx(
    row: list[str],
    uploaded_images: dict[tuple[str, str], list[bytes]],
    discipline: str,
    img_width_mm: int,
    img_height_mm: int,
    spacing_mm: int,
    add_border: bool,
) -> bytes:
    """Render one report row to DOCX bytes via the same template pipeline."""
    row_padded = (row + [""] * len(REPORT_HEADERS))[: len(REPORT_HEADERS)]
    key = (row_padded[1].strip(), row_padded[0].strip())

    single_map: dict[tuple[str, str], list[bytes]] = {}
    if key in uploaded_images:
        single_map[key] = uploaded_images[key]

    zip_bytes = generate_reports(
        [row_padded],
        single_map,
        discipline,
        img_width_mm,
        img_height_mm,
        spacing_mm,
        img_per_row=2,
        add_border=add_border,
    )

    try:
        with zipfile.ZipFile(BytesIO(zip_bytes), "r") as zf:
            docx_names = [name for name in zf.namelist() if name.lower().endswith(".docx")]
            if not docx_names:
                return b""
            return zf.read(docx_names[0])
    except (zipfile.BadZipFile, KeyError):
        return b""


def _convert_via_win32com(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    """Try converting DOCX to PDF with pywin32 Word COM."""
    try:
        import pythoncom
        import win32com.client  # type: ignore
    except Exception:
        return False, "pywin32 is not available in this environment."

    word = None
    doc = None
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path), False, True)
        doc.SaveAs(str(pdf_path), FileFormat=17)
    except Exception as exc:
        return False, f"pywin32 Word conversion failed: {exc}"
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

    if not pdf_path.exists():
        return False, "pywin32 completed but no PDF was created."
    return True, ""


def _find_powershell_executable() -> str:
    """Locate a usable PowerShell executable on Windows hosts."""
    candidates: list[str] = []

    for cmd in ("powershell", "powershell.exe", "pwsh", "pwsh.exe"):
        located = shutil.which(cmd)
        if located:
            candidates.append(located)

    system_root = os.environ.get("SystemRoot", r"C:\Windows")
    candidates.extend(
        [
            rf"{system_root}\System32\WindowsPowerShell\v1.0\powershell.exe",
            rf"{system_root}\SysWOW64\WindowsPowerShell\v1.0\powershell.exe",
            rf"{system_root}\System32\WindowsPowerShell\v1.0\pwsh.exe",
        ]
    )

    for candidate in candidates:
        try:
            if candidate and Path(candidate).exists():
                return str(Path(candidate))
        except Exception:
            continue

    return ""


def _convert_via_word_powershell(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    """Try converting DOCX to PDF with Word COM in PowerShell."""
    powershell_exe = _find_powershell_executable()
    if not powershell_exe:
        return False, "PowerShell executable not found (tried PATH + standard Windows locations)."

    ps_script = textwrap.dedent(
        """
        param(
          [Parameter(Mandatory = $true)][string]$InputPath,
          [Parameter(Mandatory = $true)][string]$OutputPath
        )

        $ErrorActionPreference = 'Stop'
        $word = $null
        $doc = $null
        try {
            $word = New-Object -ComObject Word.Application
            $word.Visible = $false
            $word.DisplayAlerts = 0
            $doc = $word.Documents.Open($InputPath, $false, $true)
            $doc.SaveAs([string]$OutputPath, 17)
        } finally {
            if ($doc -ne $null) { $doc.Close($false) }
            if ($word -ne $null) { $word.Quit() }
        }
        """
    ).strip()

    script_path = docx_path.parent / "word_to_pdf.ps1"
    script_path.write_text(ps_script, encoding="utf-8")

    try:
        result = subprocess.run(
            [
                powershell_exe,
                "-NoLogo",
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script_path),
                "-InputPath",
                str(docx_path),
                "-OutputPath",
                str(pdf_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except Exception as exc:
        return False, f"Word/PowerShell launch failed: {exc}"

    if result.returncode != 0:
        detail = (result.stderr or result.stdout or "").strip()
        return False, detail or "Word COM conversion failed."
    if not pdf_path.exists():
        return False, "Word conversion returned success but no PDF was created."
    return True, ""


def _convert_via_docx2pdf(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    """Try converting DOCX to PDF via docx2pdf (if installed)."""
    try:
        from docx2pdf import convert as docx2pdf_convert
    except Exception:
        return False, "docx2pdf is not available in this environment."

    try:
        docx2pdf_convert(str(docx_path), str(pdf_path))
    except Exception as exc:
        return False, f"docx2pdf failed: {exc}"

    if not pdf_path.exists():
        return False, "docx2pdf completed but no PDF was created."
    return True, ""


def _convert_via_libreoffice(docx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    """Try converting DOCX to PDF via LibreOffice headless mode."""
    office_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if not office_bin:
        return False, "LibreOffice is not installed."

    try:
        result = subprocess.run(
            [
                office_bin,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(docx_path.parent),
                str(docx_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=180,
        )
    except Exception as exc:
        return False, f"LibreOffice launch failed: {exc}"

    if result.returncode != 0:
        detail = (result.stderr or result.stdout or "").strip()
        return False, detail or "LibreOffice conversion failed."

    generated_pdf = docx_path.with_suffix(".pdf")
    if generated_pdf.exists() and generated_pdf != pdf_path:
        generated_pdf.replace(pdf_path)
    if not pdf_path.exists():
        return False, "LibreOffice completed but no PDF was created."
    return True, ""


def _convert_docx_to_pdf_bytes(docx_bytes: bytes) -> tuple[bytes, str]:
    """Convert DOCX bytes to PDF bytes and return (pdf_bytes, error_detail)."""
    if not docx_bytes:
        return b"", "No report document generated for preview."

    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_path = (Path(tmp_dir) / "preview.docx").resolve()
        pdf_path = (Path(tmp_dir) / "preview.pdf").resolve()
        docx_path.write_bytes(docx_bytes)

        attempts: list[str] = []
        converters = (
            _convert_via_win32com,
            _convert_via_word_powershell,
            _convert_via_docx2pdf,
            _convert_via_libreoffice,
        )
        for converter in converters:
            ok, detail = converter(docx_path, pdf_path)
            if ok and pdf_path.exists():
                return pdf_path.read_bytes(), ""
            if detail:
                attempts.append(detail)

        if not attempts:
            attempts.append("No PDF converter backend succeeded.")
        return b"", " | ".join(attempts)


def _pdf_preview_iframe_html(pdf_bytes: bytes) -> str:
    """Embed PDF bytes in an iframe for near-perfect report preview."""
    if not pdf_bytes:
        return ""

    encoded = base64.b64encode(pdf_bytes).decode("utf-8")
    return textwrap.dedent(
        f"""
        <div style="border:1px solid #d0d5dd; border-radius:10px; background:#f5f7fa; padding:12px;">
          <iframe
            src="data:application/pdf;base64,{encoded}"
            style="width:100%; height:1100px; border:1px solid #c3cad4; border-radius:8px; background:#fff;"
          ></iframe>
        </div>
        """
    ).strip()


def run_app():
    """Render the Streamlit interface."""
    set_background("bg.jpg")
    render_hero(
        title="Smart Field Reporting for Electrical and Civil Works",
        subtitle="A modern reporting system for engineers, supervisors and consultants.",
        cta_primary="Generate Reports",
        cta_secondary="Upload Site Data",
        image_path="bg.jpg",
    )
    _safe_markdown('<div id="reports-section"></div>', unsafe_allow_html=True)
    render_workwatch_header()
    _render_openai_chat_panel()

    _safe_markdown(
        """
        <style>
        div[data-testid="stRadio"] label {
            font-size: 1rem;
            font-weight: 600;
        }
        div[data-testid="stMultiSelect"] label {
            font-size: 1rem;
            font-weight: 600;
        }
        div[data-testid="stMultiSelect"] input {
            min-height: 44px;
            font-size: 0.95rem;
        }
        div[data-baseweb="select"] > div {
            min-height: 46px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.sidebar.subheader("Gallery Controls")
    img_width_mm = st.sidebar.slider(
        "Image width (mm)", min_value=50, max_value=250, value=185, step=5
    )
    img_height_mm = st.sidebar.slider(
        "Image height (mm)", min_value=50, max_value=250, value=148, step=5
    )
    st.sidebar.caption(
        "Images default to 185 mm x 148 mm each, arranged two per row with a 5 mm gap."
    )
    add_border = st.sidebar.checkbox("Add border to images", value=False)
    spacing_mm = st.sidebar.slider(
        "Gap between images (mm)", min_value=0, max_value=20, value=5, step=1
    )

    data_rows, sites, data_error = _load_sheet_context()

    if data_error is not None:  # pragma: no cover - user notification
        st.error(f"Failed to load site data: {data_error}")
        return

    container_fn = getattr(st, "container", None)
    filters_container = container_fn() if callable(container_fn) else nullcontext()
    with filters_container:
        discipline_column, selectors_column = _safe_columns((0.9, 1.8), gap="large")

        with discipline_column:
            discipline = st.radio(
                "Discipline",
                ["Civil", "Electrical"],
                key="discipline_radio",
            )

        with selectors_column:
            sites_column, dates_column = _safe_columns(2, gap="large")

            site_options = ["All Sites"] + sites if sites else ["All Sites"]
            default_sites = ["All Sites"] if site_options else []

            with sites_column:
                st.subheader("Select Sites")
                selected_sites_raw = st.multiselect(
                    "Choose sites:",
                    site_options,
                    default=default_sites,
                    key="sites_ms",
                )

            if "All Sites" in selected_sites_raw or not selected_sites_raw:
                selected_sites = sites.copy()
            else:
                selected_sites = selected_sites_raw

            available_dates = sorted(
                {
                    row[0].strip()
                    for row in data_rows
                    if not selected_sites or row[1].strip() in selected_sites
                }
            )

            date_options = ["All Dates"] + available_dates if available_dates else ["All Dates"]
            default_dates = ["All Dates"] if available_dates else []

            with dates_column:
                st.subheader("Select Dates")
                selected_dates_raw = st.multiselect(
                    "Choose dates:",
                    date_options,
                    default=default_dates,
                    key="dates_ms",
                )

            if "All Dates" in selected_dates_raw or not selected_dates_raw:
                selected_dates = available_dates
            else:
                selected_dates = selected_dates_raw

    cache = load_offline_cache()
    if cache and cache.get("rows"):
        st.info(
            "Cached offline data detected. Use the button below to sync back to the Google Sheet."
        )
        if st.button("Sync cached data to Google Sheet"):
            try:
                append_rows_to_sheet(cache.get("rows", []))
                CACHE_FILE.unlink()
                st.success("Cached data synced to Google Sheet.")
            except Exception as exc:  # pragma: no cover - user notification
                st.error(f"Sync failed: {exc}")

    filtered_rows = [
        row
        for row in data_rows
        if row[1].strip() in selected_sites and row[0].strip() in selected_dates
    ]

    site_date_pairs = sorted({(row[1].strip(), row[0].strip()) for row in filtered_rows})

    st.subheader("Preview Reports to be Generated")
    df_preview = pd.DataFrame(filtered_rows, columns=REPORT_HEADERS)
    st.dataframe(df_preview)

    st.subheader("Review Before Download")
    _safe_caption(
        "Edit report text before generation. Date and Site_Name are locked so uploaded images stay linked."
    )

    review_df = _safe_data_editor(
        df_preview,
        use_container_width=True,
        hide_index=True,
        disabled=["Date", "Site_Name"],
        key="review_editor",
    )
    review_rows = _normalized_review_rows(review_df)

    if not review_rows:
        review_rows = [
            (row + [""] * len(REPORT_HEADERS))[: len(REPORT_HEADERS)]
            for row in filtered_rows
        ]

    cabin_rows = 0
    for row in review_rows:
        sign_info = signatories_for_row(
            discipline,
            row[1],
            row[3],
            row[6],
            row[8],
            row[7],
        )
        if sign_info.get("Contractor_Name") == "Rutarindwa Olivier":
            cabin_rows += 1

    if cabin_rows:
        st.info(
            f"{cabin_rows} report(s) include cabin activities. "
            "Contractor representative is set to Rutarindwa Olivier (Civil Engineer)."
        )

    structured_from_rows = _rows_to_structured_data(review_rows)
    if st.session_state.get("_structured_origin") != "manual":
        st.session_state["structured_report_data"] = structured_from_rows
        st.session_state["_structured_origin"] = "rows"

    _safe_markdown('<div id="upload-section"></div>', unsafe_allow_html=True)

    for site, date in site_date_pairs:
        files = st.file_uploader(
            f"Upload images for {site} - {date}",
            accept_multiple_files=True,
            type=["png", "jpg", "jpeg", "webp"],
            key=f"uploader_{site}_{date}",
        )
        if files:
            key = (site.strip(), date.strip())
            st.session_state.setdefault("images", {})[key] = [f.read() for f in files]
            _safe_image(st.session_state["images"][key], width=220)

    st.subheader("Template-Size Report Preview")
    _safe_caption(
        "Template-size preview from the actual generated report document (same page size and margins as template)."
    )

    markdown_available = callable(getattr(st, "markdown", None))

    if review_rows and markdown_available:
        preview_options = list(range(len(review_rows)))
        preview_idx = _safe_selectbox(
            "Choose report to preview",
            options=preview_options,
            index=0,
            format_func=lambda idx: f"{review_rows[idx][1]} - {review_rows[idx][0]}",
            key="preview_report_select",
        )
        if preview_idx is None:
            preview_idx = 0

        preview_row = review_rows[int(preview_idx)]
        preview_docx = _build_single_report_docx(
            preview_row,
            st.session_state.get("images", {}),
            discipline,
            img_width_mm,
            img_height_mm,
            spacing_mm,
            add_border,
        )

        if preview_docx:
            with st.spinner("Rendering exact preview from template..."):
                preview_pdf, preview_pdf_error = _convert_docx_to_pdf_bytes(preview_docx)

            if preview_pdf:
                _safe_markdown(
                    _pdf_preview_iframe_html(preview_pdf),
                    unsafe_allow_html=True,
                )
                st.download_button(
                    "Download Preview PDF",
                    preview_pdf,
                    file_name="preview.pdf",
                    mime="application/pdf",
                    key=f"preview_pdf_{preview_idx}",
                )
            else:
                st.warning(
                    "Exact PDF preview is unavailable on this machine; showing best template-based preview instead."
                )
                if preview_pdf_error:
                    st.caption(f"Preview conversion detail: {preview_pdf_error}")
                _safe_markdown(
                    _render_template_preview_html(preview_docx),
                    unsafe_allow_html=True,
                )
        else:
            st.warning("Unable to render template preview for the selected row.")
    elif review_rows:
        st.info("Template preview is available in the full Streamlit UI.")
    else:
        st.info("No rows available for template preview.")

    st.subheader("Contractor Report Parser")
    enable_parser = _safe_checkbox(
        "Enable manual contractor report parsing", value=False, key="enable_parser"
    )

    if enable_parser:
        raw_report_text = st.text_area("Paste contractor report text")
        if st.button("Clean & Structure Report"):
            try:
                structured_report = clean_and_structure_report(raw_report_text)
            except (TypeError, ValueError) as exc:
                st.warning(f"Unable to structure report: {exc}")
            else:
                st.session_state["structured_report_data"] = structured_report
                st.session_state["_structured_origin"] = "manual"

    st.json(st.session_state.get("structured_report_data", structured_from_rows))

    if st.button("Generate Reports"):
        if not review_rows:
            st.warning("No data available for the selected sites and dates.")
            return

        zip_bytes = generate_reports(
            review_rows,
            st.session_state.get("images", {}),
            discipline,
            img_width_mm,
            img_height_mm,
            spacing_mm,
            img_per_row=2,
            add_border=add_border,
        )
        st.download_button("Download ZIP", zip_bytes, "reports.zip")


if __name__ == "__main__":
    run_app()

