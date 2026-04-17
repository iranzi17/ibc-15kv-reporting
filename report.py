import os
import tempfile
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional

import streamlit as st
from docxtpl import DocxTemplate, InlineImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm
from PIL import Image, ImageDraw, ImageFont, ImageOps

from config import TEMPLATE_PATH

BASE_DIR = Path(__file__).parent.resolve()
EMU_PER_MM = Mm(1).emu
JPEG_QUALITY = 92
GALLERY_RENDER_DPI = 160
RESAMPLE_LANCZOS = getattr(getattr(Image, "Resampling", Image), "LANCZOS")
TOP_SLOT_ASPECT_RATIO = 900 / 1150
BOTTOM_SLOT_ASPECT_RATIO = 1600 / 1060
LANDSCAPE_ASPECT_THRESHOLD = 1.05
PORTRAIT_ASPECT_THRESHOLD = 0.95

SIGNATORIES = {
    "Civil": {
        "Consultant_Name": "IRANZI Prince Jean Claude",
        "Consultant_Title": "Civil Engineer",
        "Consultant_Signature": "iranzi_prince_jean_claude.jpg",
        "Contractor_Name": "HABIMANA ISAAC",
        "Contractor_Title": "Electrical Engineer",
        "Contractor_Signature": "issac_habimana.jpg",
    },
    "Electrical": {
        "Consultant_Name": "Alexis IVUGIZA",
        "Consultant_Title": "Electrical Engineer",
        "Consultant_Signature": "alexis_ivugiza.jpg",
        "Contractor_Name": "HABIMANA ISAAC",
        "Contractor_Title": "Electrical Engineer",
        "Contractor_Signature": "issac_habimana.jpg",
    },
}

PLACEHOLDER_REPLACEMENTS = {
    "Reaction&amp;WayForword": "Reaction_and_WayForword",
}

EMPTY_GALLERY_MESSAGE = "No site photo was attached to the daily return."
EXTRA_GALLERY_MESSAGE = "No additional site photo was attached."
WIDE_GALLERY_MESSAGE = "No wide site photo was attached to the daily return."


def signatories_for_row(
    discipline: str,
    site_name: str,
    work: str,
    work_executed: str,
    another_work_executed: str,
    comment_on_work: str,
) -> dict[str, str]:
    """Get the fixed signatory set for the given discipline."""
    _ = (site_name, work, work_executed, another_work_executed, comment_on_work)
    return dict(SIGNATORIES.get(discipline, {}))


def safe_filename(s: str, max_len: int = 150) -> str:
    """Remove illegal filename characters and tidy whitespace."""
    import re

    s = str(s)
    s = re.sub(r'[\\/:*?"<>|]+', "-", s)
    s = re.sub(r"\s+", " ", s).strip(" .-")
    return s[:max_len]


def normalize_date(d) -> str:
    """Normalize date like '06/08/2025' -> '2025-08-06'."""
    import pandas as pd

    try:
        return pd.to_datetime(d, dayfirst=True, errors="raise").strftime("%Y-%m-%d")
    except Exception:
        return str(d).replace("/", "-").replace("\\", "-")


def format_date_title(d: str) -> str:
    """Return dd.MM.YYYY for filenames like 04.08.2025."""
    import pandas as pd

    try:
        return pd.to_datetime(d, dayfirst=True, errors="raise").strftime("%d.%m.%Y")
    except Exception:
        return str(d).replace("/", ".").replace("-", ".")


def resolve_asset(name: Optional[str]) -> Optional[str]:
    """Find an asset whether it's in ./ or ./signatures/, with or without extension."""
    if not name:
        return None
    p = (BASE_DIR / name).resolve()
    stem = p.with_suffix("").name
    if p.parent != BASE_DIR:
        search_dirs = [p.parent]
    else:
        search_dirs = [BASE_DIR / "signatures", BASE_DIR]
    exts = ["", ".png", ".jpg", ".jpeg", ".webp"]
    for d in search_dirs:
        for ext in exts:
            candidate = d / f"{stem}{ext}"
            if candidate.exists():
                return str(candidate)
    return None


def _create_sanitized_template_copy(template_path: str) -> str:
    """Return a temporary copy of the template with problematic placeholders normalised."""

    with zipfile.ZipFile(template_path, "r") as src, tempfile.NamedTemporaryFile(
        delete=False, suffix=".docx"
    ) as tmp:
        with zipfile.ZipFile(tmp, "w") as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename.endswith(".xml") and PLACEHOLDER_REPLACEMENTS:
                    try:
                        text = data.decode("utf-8")
                    except UnicodeDecodeError:
                        pass
                    else:
                        for old, new in PLACEHOLDER_REPLACEMENTS.items():
                            text = text.replace(old, new)
                        data = text.encode("utf-8")
                dst.writestr(item, data)
    return tmp.name


def _mm_to_twips(mm_value: float) -> int:
    """Convert millimetres to Word twips (1/20th of a point)."""

    twips = mm_value * 1440 / 25.4
    return max(0, int(round(twips)))


def _set_cell_margin(cell, side: str, value_mm: float) -> None:
    """Set an individual cell margin in millimetres."""

    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    margin = tc_mar.find(qn(f"w:{side}"))
    if margin is None:
        margin = OxmlElement(f"w:{side}")
        tc_mar.append(margin)
    margin.set(qn("w:w"), str(_mm_to_twips(value_mm)))
    margin.set(qn("w:type"), "dxa")


def _mm_to_pixels(mm_value: float, *, dpi: int = GALLERY_RENDER_DPI) -> int:
    """Convert millimetres to pixels for pre-sizing gallery images."""

    pixels = float(mm_value) * dpi / 25.4
    return max(1, int(round(pixels)))


def _gallery_slot_size_px(width_mm: float, height_mm: float | None) -> tuple[int, int]:
    """Return a pixel size matching the selected physical gallery slot."""

    safe_width_mm = max(1.0, float(width_mm))
    safe_height_mm = max(1.0, float(height_mm or width_mm))
    return _mm_to_pixels(safe_width_mm), _mm_to_pixels(safe_height_mm)


def _load_gallery_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """Load a readable system font for placeholder images."""

    candidates = [
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/Library/Fonts/Arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            try:
                return ImageFont.truetype(str(candidate), size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def _wrap_text_by_pixels(
    text: str,
    draw: ImageDraw.ImageDraw,
    font: ImageFont.ImageFont,
    max_width: int,
) -> str:
    """Wrap placeholder text to a target pixel width."""

    words = str(text or "").split()
    if not words:
        return ""

    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        bbox = draw.textbbox((0, 0), candidate, font=font)
        width = bbox[2] - bbox[0]
        if width <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return "\n".join(lines)


def _placeholder_gallery_image_bytes(size: tuple[int, int], message: str) -> bytes:
    """Return a centered JPEG placeholder sized for the report gallery slot."""

    width, height = size
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    margin = max(20, min(width, height) // 30)
    border_width = max(2, min(width, height) // 250)
    draw.rectangle(
        (margin, margin, width - margin, height - margin),
        outline=(218, 218, 218),
        width=border_width,
    )

    font_size = max(24, min(width, height) // 20)
    font = _load_gallery_font(font_size)
    wrapped = _wrap_text_by_pixels(message, draw, font, max(1, width - (margin * 4)))
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=8, align="center")
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = (width - text_width) / 2
    y = (height - text_height) / 2

    draw.multiline_text(
        (x, y),
        wrapped,
        font=font,
        fill=(120, 120, 120),
        spacing=8,
        align="center",
    )

    buffer = BytesIO()
    image.save(buffer, format="JPEG", quality=JPEG_QUALITY, optimize=True)
    return buffer.getvalue()


def _prepared_gallery_image_bytes(
    image_bytes: bytes | None,
    *,
    size: tuple[int, int],
    missing_message: str,
    failure_message: str,
) -> bytes:
    """Return a normalized gallery JPEG or a placeholder when the input is unusable."""

    if not image_bytes:
        return _placeholder_gallery_image_bytes(size, missing_message)

    try:
        with Image.open(BytesIO(image_bytes)) as img:
            image = ImageOps.exif_transpose(img).convert("RGB")
            fitted = ImageOps.fit(
                image,
                size,
                method=RESAMPLE_LANCZOS,
                centering=(0.5, 0.5),
            )
    except Exception:
        return _placeholder_gallery_image_bytes(size, failure_message)

    buffer = BytesIO()
    fitted.save(buffer, format="JPEG", quality=JPEG_QUALITY, optimize=True)
    return buffer.getvalue()


def _gallery_placeholder_message(index: int) -> str:
    """Return a user-facing message for a missing gallery photo slot."""

    if index <= 1:
        return EMPTY_GALLERY_MESSAGE
    return EXTRA_GALLERY_MESSAGE


def _image_size_from_bytes(image_bytes: bytes | None) -> tuple[int, int] | None:
    """Return image dimensions from raw bytes, or None when unreadable."""

    if not image_bytes:
        return None
    try:
        with Image.open(BytesIO(image_bytes)) as img:
            width, height = img.size
    except Exception:
        return None

    if width <= 0 or height <= 0:
        return None
    return width, height


def _image_aspect_ratio(image_bytes: bytes | None) -> float | None:
    """Return width/height for an image, or None when unreadable."""

    size = _image_size_from_bytes(image_bytes)
    if not size:
        return None
    width, height = size
    return width / height


def _gallery_layout_geometry(
    gallery_width_mm: float,
    wide_photo_height_mm: float | None,
    spacing_mm: float,
) -> dict[str, int]:
    """Return pixel geometry for the editorial gallery layout."""

    total_width_px = _mm_to_pixels(max(1.0, float(gallery_width_mm)))
    gap_px = _mm_to_pixels(max(0.0, float(spacing_mm)))
    top_slot_width_px = max(1, (total_width_px - gap_px) // 2)
    top_slot_height_px = max(1, int(round(top_slot_width_px / TOP_SLOT_ASPECT_RATIO)))

    if wide_photo_height_mm:
        bottom_slot_height_px = _mm_to_pixels(max(1.0, float(wide_photo_height_mm)))
    else:
        bottom_slot_height_px = max(1, int(round(total_width_px / BOTTOM_SLOT_ASPECT_RATIO)))

    return {
        "total_width_px": total_width_px,
        "gap_px": gap_px,
        "top_slot_width_px": top_slot_width_px,
        "top_slot_height_px": top_slot_height_px,
        "bottom_slot_height_px": bottom_slot_height_px,
    }


def _gallery_slot_boxes(geometry: dict[str, int]) -> dict[str, tuple[int, int, int, int]]:
    """Return slot rectangles as (left, top, width, height)."""

    top_width = geometry["top_slot_width_px"]
    top_height = geometry["top_slot_height_px"]
    total_width = geometry["total_width_px"]
    gap_px = geometry["gap_px"]
    bottom_height = geometry["bottom_slot_height_px"]

    return {
        "top_left": (0, 0, top_width, top_height),
        "top_right": (top_width + gap_px, 0, top_width, top_height),
        "bottom": (0, top_height + gap_px, total_width, bottom_height),
        "wide": (0, 0, total_width, bottom_height),
    }


def _gallery_page_groups(image_bytes_list: List[bytes], *, max_per_page: int = 3) -> List[List[bytes]]:
    """Split uploaded site photos into fixed-size gallery pages."""

    images = list(image_bytes_list or [])
    return [images[index : index + max_per_page] for index in range(0, len(images), max_per_page)]


def _draw_gallery_border(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
) -> None:
    """Draw a subtle slot border around one gallery image."""

    left, top, width, height = box
    border_width = max(2, min(width, height) // 250)
    draw.rectangle(
        (left, top, left + width - 1, top + height - 1),
        outline=(218, 218, 218),
        width=border_width,
    )


def _paste_gallery_slot(
    canvas: Image.Image,
    draw: ImageDraw.ImageDraw,
    *,
    image_bytes: bytes | None,
    box: tuple[int, int, int, int],
    missing_message: str,
    failure_message: str,
    add_border: bool,
    use_placeholder: bool,
) -> None:
    """Paste one prepared gallery slot into the composed page canvas."""

    if image_bytes is None and not use_placeholder:
        return

    left, top, width, height = box
    slot_bytes = _prepared_gallery_image_bytes(
        image_bytes,
        size=(width, height),
        missing_message=missing_message,
        failure_message=failure_message,
    )
    with Image.open(BytesIO(slot_bytes)) as slot_image:
        canvas.paste(slot_image.convert("RGB"), (left, top))

    if add_border:
        _draw_gallery_border(draw, box)


def _gallery_layout_name(page_images: List[bytes]) -> str:
    """Choose the editorial layout for one gallery page."""

    count = len(page_images)
    if count <= 0:
        return "empty"
    if count >= 3:
        return "two_top_one_bottom"

    first_aspect = _image_aspect_ratio(page_images[0])
    second_aspect = _image_aspect_ratio(page_images[1]) if count > 1 else None

    if count == 2:
        if (
            first_aspect is not None
            and second_aspect is not None
            and first_aspect <= PORTRAIT_ASPECT_THRESHOLD
            and second_aspect >= LANDSCAPE_ASPECT_THRESHOLD
        ):
            return "portrait_plus_wide"
        return "two_top"

    if first_aspect is not None and first_aspect <= PORTRAIT_ASPECT_THRESHOLD:
        return "single_portrait"
    return "single_wide"


def _compose_gallery_page_bytes(
    page_images: List[bytes],
    *,
    gallery_width_mm: float,
    wide_photo_height_mm: float | None,
    spacing_mm: float,
    add_border: bool,
    show_photo_placeholders: bool,
) -> bytes | None:
    """Compose one report gallery page as a single JPEG image."""

    layout_name = _gallery_layout_name(page_images)
    if layout_name == "empty" and not show_photo_placeholders:
        return None

    geometry = _gallery_layout_geometry(gallery_width_mm, wide_photo_height_mm, spacing_mm)
    boxes = _gallery_slot_boxes(geometry)
    total_width_px = geometry["total_width_px"]
    gap_px = geometry["gap_px"]
    top_height_px = geometry["top_slot_height_px"]
    bottom_height_px = geometry["bottom_slot_height_px"]

    if layout_name in {"two_top_one_bottom", "portrait_plus_wide"}:
        canvas_height_px = top_height_px + gap_px + bottom_height_px
    elif layout_name == "two_top":
        canvas_height_px = top_height_px
    else:
        canvas_height_px = bottom_height_px if layout_name == "single_wide" else top_height_px

    canvas = Image.new("RGB", (total_width_px, canvas_height_px), "white")
    draw = ImageDraw.Draw(canvas)

    if layout_name == "empty":
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=None,
            box=boxes["wide"],
            missing_message=EMPTY_GALLERY_MESSAGE,
            failure_message="Unable to prepare the gallery placeholder.",
            add_border=add_border,
            use_placeholder=True,
        )
    elif layout_name == "two_top_one_bottom":
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[0],
            box=boxes["top_left"],
            missing_message=_gallery_placeholder_message(1),
            failure_message="Unable to process site photo 1.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[1],
            box=boxes["top_right"],
            missing_message=_gallery_placeholder_message(2),
            failure_message="Unable to process site photo 2.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[2],
            box=boxes["bottom"],
            missing_message=WIDE_GALLERY_MESSAGE,
            failure_message="Unable to process site photo 3.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
    elif layout_name == "portrait_plus_wide":
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[0],
            box=boxes["top_left"],
            missing_message=_gallery_placeholder_message(1),
            failure_message="Unable to process site photo 1.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[1],
            box=boxes["bottom"],
            missing_message=WIDE_GALLERY_MESSAGE,
            failure_message="Unable to process site photo 2.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
    elif layout_name == "two_top":
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[0],
            box=boxes["top_left"],
            missing_message=_gallery_placeholder_message(1),
            failure_message="Unable to process site photo 1.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[1],
            box=boxes["top_right"],
            missing_message=_gallery_placeholder_message(2),
            failure_message="Unable to process site photo 2.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
    elif layout_name == "single_portrait":
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[0],
            box=boxes["top_left"],
            missing_message=_gallery_placeholder_message(1),
            failure_message="Unable to process site photo 1.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )
    else:
        _paste_gallery_slot(
            canvas,
            draw,
            image_bytes=page_images[0],
            box=boxes["wide"],
            missing_message=EMPTY_GALLERY_MESSAGE,
            failure_message="Unable to process site photo 1.",
            add_border=add_border,
            use_placeholder=show_photo_placeholders,
        )

    buffer = BytesIO()
    canvas.save(buffer, format="JPEG", quality=JPEG_QUALITY, optimize=True)
    return buffer.getvalue()


def _apply_cell_spacing(cell, spacing_mm: float, column_index: int, total_columns: int) -> tuple[float, float]:
    """Apply asymmetric margins that create an even gap between images."""

    spacing_mm = max(0.0, float(spacing_mm))
    top_bottom = spacing_mm
    if total_columns <= 1:
        left = right = spacing_mm
    else:
        inner_gap = spacing_mm / 2.0
        left = spacing_mm if column_index == 0 else inner_gap
        right = spacing_mm if column_index == total_columns - 1 else inner_gap

    for side, value in ("top", top_bottom), ("bottom", top_bottom), ("left", left), ("right", right):
        _set_cell_margin(cell, side, value)

    return left, right


def generate_reports(
    filtered_rows: List[List[str]],
    uploaded_image_mapping: Dict[tuple, List[bytes]],
    discipline: str,
    img_width_mm: int,
    img_height_mm: int,
    spacing_mm: int,
    img_per_row: int = 2,
    add_border: bool = False,
    show_photo_placeholders: bool = True,
    template_path: str = TEMPLATE_PATH,
) -> bytes:
    """Create a ZIP archive of rendered DOCX reports."""
    zip_buffer = BytesIO()
    sanitized_template = _create_sanitized_template_copy(template_path)
    gallery_width_mm = max(1.0, float(img_width_mm))
    wide_photo_height_mm = float(img_height_mm) if img_height_mm else None

    try:
        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            used_names: Dict[str, int] = {}
            for row in filtered_rows:
                (
                    date,
                    site_name,
                    district,
                    work,
                    human_resources,
                    supply,
                    work_executed,
                    comment_on_work,
                    another_work_executed,
                    comment_on_hse,
                    consultant_recommandation,
                    non_compliant_work,
                    reaction_way_forward,
                    challenges,
                ) = (row + [""] * 14)[:14]
                site_name = site_name.strip()
                date = date.strip()

                tpl = DocxTemplate(sanitized_template)

                required = {
                    "Consultant_Name",
                    "Consultant_Title",
                    "Consultant_Signature",
                    "Contractor_Name",
                    "Contractor_Title",
                    "Contractor_Signature",
                }
                placeholders = tpl.get_undeclared_template_variables({})
                missing = required - placeholders
                if missing:
                    st.warning(
                        "Template is missing placeholders: " + ", ".join(sorted(missing))
                    )

                image_bytes = uploaded_image_mapping.get((site_name, date), []) or []
                gallery_groups = _gallery_page_groups(image_bytes)
                if not gallery_groups and show_photo_placeholders:
                    gallery_groups = [[]]

                images_subdoc = tpl.new_subdoc()
                for index, gallery_group in enumerate(gallery_groups):
                    gallery_page_bytes = _compose_gallery_page_bytes(
                        gallery_group,
                        gallery_width_mm=gallery_width_mm,
                        wide_photo_height_mm=wide_photo_height_mm,
                        spacing_mm=spacing_mm,
                        add_border=add_border,
                        show_photo_placeholders=show_photo_placeholders,
                    )
                    if not gallery_page_bytes:
                        continue

                    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_img:
                        tmp_img.write(gallery_page_bytes)
                        tmp_img.flush()
                        gallery_path = tmp_img.name

                    paragraph = images_subdoc.add_paragraph()
                    paragraph.alignment = 1
                    paragraph.add_run().add_picture(gallery_path, width=Mm(gallery_width_mm))
                    try:
                        os.remove(gallery_path)
                    except OSError:
                        pass

                    if index != len(gallery_groups) - 1:
                        images_subdoc.add_page_break()

                sign_info = signatories_for_row(
                    discipline,
                    site_name,
                    work,
                    work_executed,
                    another_work_executed,
                    comment_on_work,
                )
                cons_sig_path = resolve_asset(sign_info.get("Consultant_Signature"))
                cont_sig_path = resolve_asset(sign_info.get("Contractor_Signature"))
                cons_sig_img = (
                    InlineImage(tpl, cons_sig_path, width=Mm(30)) if cons_sig_path else ""
                )
                cont_sig_img = (
                    InlineImage(tpl, cont_sig_path, width=Mm(30)) if cont_sig_path else ""
                )

                ctx = {
                    "Date": date,
                    "Site_Name": site_name,
                    "District": district,
                    "Work": work,
                    "Human_Resources": human_resources,
                    "Supply": supply,
                    "Work_Executed": work_executed,
                    "Comment_on_work": comment_on_work,
                    "Another_Work_Executed": another_work_executed,
                    "Comment_on_HSE": comment_on_hse,
                    "Consultant_Recommandation": consultant_recommandation,
                    "Non_Compliant_work": non_compliant_work,
                    "Reaction_and_WayForword": reaction_way_forward,
                    "challenges": challenges,
                    "Consultant_Name": sign_info.get("Consultant_Name", ""),
                    "Consultant_Title": sign_info.get("Consultant_Title", ""),
                    "Contractor_Name": sign_info.get("Contractor_Name", ""),
                    "Contractor_Title": sign_info.get("Contractor_Title", ""),
                    "Consultant_Signature": cons_sig_img,
                    "Contractor_Signature": cont_sig_img,
                    "Images": images_subdoc,
                }

                # Backwards compatibility for templates that still use the unsanitised placeholder.
                ctx["Reaction&WayForword"] = reaction_way_forward

                tpl.render(ctx)

                base_filename = safe_filename(
                    "_".join(filter(None, [site_name or "report", format_date_title(date)]))
                )
                if not base_filename:
                    base_filename = "report"
                count = used_names.get(base_filename, 0) + 1
                used_names[base_filename] = count
                filename = base_filename if count == 1 else f"{base_filename}_{count}"
                if not filename.lower().endswith(".docx"):
                    filename = f"{filename}.docx"

                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_doc:
                    tpl.save(tmp_doc.name)
                    tmp_path = tmp_doc.name
                with open(tmp_path, "rb") as fh:
                    zipf.writestr(filename, fh.read())
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass

    finally:
        try:
            os.remove(sanitized_template)
        except OSError:
            pass
    zip_buffer.seek(0)
    return zip_buffer.getvalue()
