from io import BytesIO
import zipfile

from docx import Document
from PIL import Image
import pytest

import report


def _png_bytes(width: int, height: int, color: tuple[int, int, int] = (255, 0, 0)) -> bytes:
    image = Image.new("RGB", (width, height), color)
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


SQUARE_PNG = _png_bytes(200, 200)
LANDSCAPE_PNG = _png_bytes(640, 320)
PORTRAIT_PNG = _png_bytes(320, 640)


def test_safe_filename():
    assert report.safe_filename('inv?lid:name') == 'inv-lid-name'


def test_format_date_title():
    assert report.format_date_title('06/08/2025') == '06.08.2025'


def _empty_row(site: str, date: str) -> list[str]:
    return [date, site] + [""] * 12


def test_signatories_for_row_keeps_habimana_isaac_for_cabin_activity():
    sign_info = report.signatories_for_row(
        "Civil",
        "KIBAGABAGA SMART CABIN",
        "Cabin rehabilitation",
        "Executed cabin works",
        "",
        "Progress on cabin works",
    )

    assert sign_info["Contractor_Name"] == "HABIMANA ISAAC"
    assert sign_info["Contractor_Title"] == "Electrical Engineer"
    assert sign_info["Contractor_Signature"] == "issac_habimana.jpg"


def test_generate_reports_returns_zip():
    rows = [_empty_row("Site A", "2025-08-06")]
    uploaded = {}
    data = report.generate_reports(rows, uploaded, 'Civil', 70, 60, 2, 2, False)
    with zipfile.ZipFile(BytesIO(data)) as zf:
        names = zf.namelist()
    assert len(names) == 1
    assert names[0].startswith('Site A')


def test_generate_reports_unique_docx_names_for_duplicates():
    rows = [
        _empty_row("Site A", "2025-08-06"),
        _empty_row("Site A", "2025-08-06"),
        _empty_row("", ""),
        _empty_row("", ""),
    ]
    data = report.generate_reports(rows, {}, "Civil", 70, 60, 2, 2, False)
    with zipfile.ZipFile(BytesIO(data)) as zf:
        names = sorted(zf.namelist())

    assert names == [
        "Site A_08.06.2025.docx",
        "Site A_08.06.2025_2.docx",
        "report.docx",
        "report_2.docx",
    ]


def test_prepared_gallery_image_bytes_resizes_to_slot():
    prepared = report._prepared_gallery_image_bytes(
        LANDSCAPE_PNG,
        size=(320, 180),
        missing_message="missing",
        failure_message="failure",
    )

    with Image.open(BytesIO(prepared)) as image:
        assert image.size == (320, 180)


def test_prepared_gallery_image_bytes_falls_back_to_placeholder_for_invalid_data():
    prepared = report._prepared_gallery_image_bytes(
        b"not-a-real-image",
        size=(240, 160),
        missing_message="missing",
        failure_message="failure",
    )

    with Image.open(BytesIO(prepared)) as image:
        assert image.size == (240, 160)


def test_gallery_layout_name_prefers_portrait_plus_wide_for_mixed_pair():
    assert report._gallery_layout_name([PORTRAIT_PNG, LANDSCAPE_PNG]) == "portrait_plus_wide"


def test_compose_gallery_page_bytes_uses_full_gallery_width():
    page = report._compose_gallery_page_bytes(
        [PORTRAIT_PNG, LANDSCAPE_PNG],
        captions=["Pole replacement in progress.", "Wide trench section ready for cable laying."],
        gallery_width_mm=185,
        wide_photo_height_mm=120,
        spacing_mm=5,
        add_border=False,
        show_photo_placeholders=False,
    )

    geometry = report._gallery_layout_geometry(185, 120, 5)
    expected_height = geometry["top_slot_height_px"] + geometry["gap_px"] + geometry["bottom_slot_height_px"]

    with Image.open(BytesIO(page)) as image:
        assert image.size == (geometry["total_width_px"], expected_height)


def test_generate_reports_accepts_image_caption_mapping():
    rows = [_empty_row("Site A", "2025-08-06")]
    uploaded = {("Site A", "2025-08-06"): [PORTRAIT_PNG, LANDSCAPE_PNG]}
    captions = {("Site A", "2025-08-06"): ["Workers preparing pole access.", "Wide view of overhead line activity."]}

    data = report.generate_reports(
        rows,
        uploaded,
        "Civil",
        185,
        120,
        5,
        img_per_row=2,
        add_border=False,
        show_photo_placeholders=False,
        image_caption_mapping=captions,
    )

    with zipfile.ZipFile(BytesIO(data)) as zf:
        doc_bytes = zf.read(zf.namelist()[0])

    document = Document(BytesIO(doc_bytes))
    assert len(document.inline_shapes) == 3


def test_generate_reports_adds_placeholder_images_when_enabled():
    rows = [_empty_row("Site A", "2025-08-06")]
    data = report.generate_reports(
        rows,
        {},
        "Civil",
        70,
        60,
        2,
        img_per_row=2,
        add_border=False,
        show_photo_placeholders=True,
    )
    with zipfile.ZipFile(BytesIO(data)) as zf:
        doc_bytes = zf.read(zf.namelist()[0])

    document = Document(BytesIO(doc_bytes))

    assert len(document.inline_shapes) == 3


def test_generate_reports_skips_placeholder_images_when_disabled():
    rows = [_empty_row("Site A", "2025-08-06")]
    data = report.generate_reports(
        rows,
        {},
        "Civil",
        70,
        60,
        2,
        img_per_row=2,
        add_border=False,
        show_photo_placeholders=False,
    )
    with zipfile.ZipFile(BytesIO(data)) as zf:
        doc_bytes = zf.read(zf.namelist()[0])

    document = Document(BytesIO(doc_bytes))

    assert len(document.inline_shapes) == 2


def test_generate_reports_uses_one_gallery_collage_for_three_photos():
    rows = [_empty_row("Site A", "2025-08-06")]
    uploaded = {("Site A", "2025-08-06"): [PORTRAIT_PNG, LANDSCAPE_PNG, SQUARE_PNG]}
    width_mm = 185
    height_mm = 120
    spacing_mm = 5
    data = report.generate_reports(
        rows,
        uploaded,
        "Civil",
        width_mm,
        height_mm,
        spacing_mm,
        img_per_row=1,
        add_border=False,
    )
    with zipfile.ZipFile(BytesIO(data)) as zf:
        docx_name = zf.namelist()[0]
        doc_bytes = zf.read(docx_name)

    document = Document(BytesIO(doc_bytes))

    assert len(document.inline_shapes) == 3


def test_generate_reports_uses_two_collage_pages_for_four_photos():
    rows = [_empty_row("Site A", "2025-08-06")]
    uploaded = {("Site A", "2025-08-06"): [PORTRAIT_PNG, LANDSCAPE_PNG, SQUARE_PNG, LANDSCAPE_PNG]}
    data = report.generate_reports(
        rows,
        uploaded,
        "Civil",
        185,
        120,
        5,
        img_per_row=2,
        add_border=False,
        show_photo_placeholders=False,
    )
    with zipfile.ZipFile(BytesIO(data)) as zf:
        doc_bytes = zf.read(zf.namelist()[0])

    document = Document(BytesIO(doc_bytes))

    assert len(document.inline_shapes) == 4

